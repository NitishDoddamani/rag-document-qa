import pdfplumber
import pytesseract
import io
import re
from typing import List
from PIL import Image
from docx import Document

# Set tesseract path (Windows)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF — handles both text-based and image-based PDFs"""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()

                # If page has text content
                if page_text and len(page_text.strip()) > 30:
                    text += f"\n\n[PAGE {page_num}]\n{page_text.strip()}\n"
                else:
                    # Page is image-based — use OCR
                    print(f"Page {page_num} has no text, attempting OCR...")
                    try:
                        page_image = page.to_image(resolution=200).original
                        ocr_text = pytesseract.image_to_string(page_image)
                        if ocr_text.strip():
                            text += f"\n\n[PAGE {page_num} - OCR]\n{ocr_text.strip()}\n"
                        else:
                            text += f"\n\n[PAGE {page_num}]\n[Image content - no extractable text]\n"
                    except Exception as ocr_err:
                        print(f"OCR failed on page {page_num}: {ocr_err}")

    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""

    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from Word documents"""
    text = ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        for para_num, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text += para.text.strip() + "\n\n"

        # Also extract tables from word doc
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text += row_text + "\n"

    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""

    return text.strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Universal text extractor — detects file type and routes accordingly
    """
    filename_lower = filename.lower()

    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
        return extract_text_from_docx(file_bytes)
    else:
        # Try PDF as fallback
        return extract_text_from_pdf(file_bytes)


def chunk_text(
    text: str,
    max_chunk_words: int = 350,
    min_chunk_words: int = 80,
    overlap_words: int = 120
) -> List[str]:
    """
    Generalized semantic chunking — works for any document type
    """
    if not text:
        return []

    # Clean up
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)

    # Generalized split pattern — works for any document
    split_pattern = r'(?=\n{2,}|\[PAGE \d+\])'
    candidates = re.split(split_pattern, text.strip())

    paragraphs = []
    for cand in candidates:
        cleaned = cand.strip()
        if cleaned and len(cleaned.split()) >= 5:
            paragraphs.append(cleaned)

    chunks = []
    current_chunk = ""
    current_word_count = 0

    for para in paragraphs:
        para_words = para.split()
        para_word_count = len(para_words)

        if current_word_count + para_word_count > max_chunk_words and current_chunk:
            chunks.append(current_chunk.strip())
            if para_word_count > max_chunk_words * 0.6:
                overlap_start = max(0, para_word_count - overlap_words)
                current_chunk = " ".join(para_words[overlap_start:])
                current_word_count = para_word_count - overlap_start
            else:
                current_chunk = para
                current_word_count = para_word_count
        else:
            separator = "\n\n" if current_chunk else ""
            current_chunk += separator + para
            current_word_count += para_word_count

    if current_chunk:
        chunks.append(current_chunk.strip())

    # Merge small chunks
    final_chunks = []
    i = 0
    while i < len(chunks):
        chunk = chunks[i]
        word_count = len(chunk.split())
        if word_count < min_chunk_words and i + 1 < len(chunks):
            merged = chunk + "\n\n" + chunks[i + 1]
            if len(merged.split()) <= max_chunk_words:
                final_chunks.append(merged)
                i += 2
                continue
        final_chunks.append(chunk)
        i += 1

    return final_chunks